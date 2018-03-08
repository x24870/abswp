#Open a text file with list of guest names
#And generate a .docx file with invatation information

import docx

with open('guests.txt', 'r') as file:
    doc = docx.Document()
    
    for idx, guest in enumerate(guests):
        doc.add_heading('Invatation', )
        doc.add_paragraph('Hi, {} I sincerely invite you to my party.'.format(guest))
        doc.paragraphs[1+idx*2].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        doc.paragraphs[1+idx*2].style = 'Quote'
        
    doc.save('Invatation.docx')